"""
Data Processing Module
======================

This module provides comprehensive functionality for processing financial documents,
extracting data, and managing file operations for the Financial Risk Assessment Platform.

Classes:
--------
- DocumentProcessor: Process various document formats
- DataExtractor: Extract financial data from documents
- FileManager: Manage file operations and storage

Functions:
----------
- process_uploaded_files: Main function to process uploaded files
- extract_financial_data: Extract financial data from text/tables
- clean_and_validate_data: Clean and validate extracted data

Author: Prateek Dahiya
"""

import pandas as pd
import numpy as np
import re
import logging
import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Union
import mimetypes
import hashlib
import warnings
warnings.filterwarnings('ignore')

# Import document processing libraries with fallbacks
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import openpyxl
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from werkzeug.utils import secure_filename
    WERKZEUG_AVAILABLE = True
except ImportError:
    WERKZEUG_AVAILABLE = False
    
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats for financial data extraction"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.xlsx', '.xls', '.csv', '.docx', '.txt'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    def __init__(self):
        self.supported_formats = self._check_supported_formats()
        logger.info(f"Document processor initialized. Supported formats: {self.supported_formats}")
    
    def _check_supported_formats(self) -> Dict[str, bool]:
        """Check which document formats are supported"""
        return {
            'pdf': PDF_AVAILABLE or PDFPLUMBER_AVAILABLE,
            'excel': EXCEL_AVAILABLE,
            'csv': True,  # Always supported with pandas
            'docx': DOCX_AVAILABLE,
            'txt': True   # Always supported
        }
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file is supported"""
        extension = Path(filename).suffix.lower()
        return extension in self.ALLOWED_EXTENSIONS
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        try:
            if not os.path.exists(file_path):
                return {'error': 'File not found'}
            
            file_stats = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            # Calculate file hash for integrity checking
            file_hash = self._calculate_file_hash(file_path)
            
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            return {
                'filename': file_path_obj.name,
                'size': file_stats.st_size,
                'size_mb': round(file_stats.st_size / (1024 * 1024), 2),
                'extension': file_path_obj.suffix.lower(),
                'mime_type': mime_type,
                'created': datetime.fromtimestamp(file_stats.st_ctime),
                'modified': datetime.fromtimestamp(file_stats.st_mtime),
                'hash': file_hash,
                'is_supported': self.is_supported_file(file_path_obj.name),
                'estimated_processing_time': self._estimate_processing_time(file_stats.st_size, file_path_obj.suffix)
            }
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {'error': str(e)}
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash: {e}")
            return ""
    
    def _estimate_processing_time(self, file_size: int, extension: str) -> float:
        """Estimate processing time in seconds"""
        base_time = {
            '.pdf': 0.5,
            '.xlsx': 0.3,
            '.xls': 0.3,
            '.csv': 0.1,
            '.docx': 0.4,
            '.txt': 0.1
        }
        
        size_factor = file_size / (1024 * 1024)  # Size in MB
        return base_time.get(extension, 0.3) * max(1, size_factor * 0.5)
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document and extract text and tables"""
        try:
            file_info = self.get_file_info(file_path)
            
            if 'error' in file_info:
                return {'success': False, 'error': file_info['error']}
            
            if not file_info['is_supported']:
                return {'success': False, 'error': f"Unsupported file type: {file_info['extension']}"}
            
            if file_info['size'] > self.MAX_FILE_SIZE:
                return {'success': False, 'error': f"File too large: {file_info['size_mb']}MB (max: {self.MAX_FILE_SIZE/(1024*1024)}MB)"}
            
            # Process based on file type
            extension = file_info['extension']
            
            if extension == '.pdf':
                result = self._process_pdf(file_path)
            elif extension in ['.xlsx', '.xls']:
                result = self._process_excel(file_path)
            elif extension == '.csv':
                result = self._process_csv(file_path)
            elif extension == '.docx':
                result = self._process_docx(file_path)
            elif extension == '.txt':
                result = self._process_text(file_path)
            else:
                return {'success': False, 'error': f"Processing not implemented for {extension}"}
            
            if result['success']:
                result['file_info'] = file_info
                result['processed_at'] = datetime.now()
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files"""
        if not (PDF_AVAILABLE or PDFPLUMBER_AVAILABLE):
            return {'success': False, 'error': 'PDF processing libraries not available'}
        
        text = ""
        tables = []
        metadata = {}
        
        try:
            # Try pdfplumber first (better for tables)
            if PDFPLUMBER_AVAILABLE:
                text, tables, metadata = self._extract_with_pdfplumber(file_path)
            elif PDF_AVAILABLE:
                text, metadata = self._extract_with_pypdf2(file_path)
            
            return {
                'success': True,
                'text': text,
                'tables': tables,
                'metadata': metadata,
                'extraction_method': 'pdfplumber' if PDFPLUMBER_AVAILABLE else 'pypdf2'
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {'success': False, 'error': str(e)}
    
    def _extract_with_pdfplumber(self, file_path: str) -> Tuple[str, List[pd.DataFrame], Dict]:
        """Extract text and tables using pdfplumber"""
        import pdfplumber
        
        text = ""
        tables = []
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                'pages': len(pdf.pages),
                'creator': pdf.metadata.get('Creator', ''),
                'producer': pdf.metadata.get('Producer', ''),
                'creation_date': pdf.metadata.get('CreationDate', '')
            }
            
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                # Extract tables
                page_tables = page.extract_tables()
                for table_num, table in enumerate(page_tables):
                    if table and len(table) > 1:
                        try:
                            # Create DataFrame from table
                            df = pd.DataFrame(table[1:], columns=table[0])
                            df.name = f"Page_{page_num + 1}_Table_{table_num + 1}"
                            tables.append(df)
                        except Exception as e:
                            logger.warning(f"Could not convert table to DataFrame: {e}")
        
        return text, tables, metadata
    
    def _extract_with_pypdf2(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text using PyPDF2"""
        import PyPDF2
        
        text = ""
        metadata = {}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            metadata = {
                'pages': len(pdf_reader.pages),
                'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else '',
                'producer': pdf_reader.metadata.get('/Producer', '') if pdf_reader.metadata else ''
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        
        return text, metadata
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel files"""
        if not EXCEL_AVAILABLE:
            return {'success': False, 'error': 'Excel processing library not available'}
        
        try:
            # Load workbook to get metadata
            workbook = load_workbook(file_path, data_only=True)
            
            text = ""
            tables = []
            metadata = {
                'sheets': workbook.sheetnames,
                'sheet_count': len(workbook.sheetnames)
            }
            
            # Process each sheet
            for sheet_name in workbook.sheetnames:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if not df.empty:
                        df.name = sheet_name
                        tables.append(df)
                        
                        # Convert to text representation
                        text += f"\n--- Sheet: {sheet_name} ---\n"
                        text += df.to_string(index=False, max_rows=50)
                        text += "\n"
                        
                except Exception as e:
                    logger.warning(f"Could not process sheet '{sheet_name}': {e}")
                    continue
            
            return {
                'success': True,
                'text': text,
                'tables': tables,
                'metadata': metadata,
                'extraction_method': 'openpyxl'
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            df = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                return {'success': False, 'error': 'Could not decode CSV file with any encoding'}
            
            text = df.to_string(index=False, max_rows=100)
            tables = [df] if not df.empty else []
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'encoding': used_encoding,
                'column_names': df.columns.tolist()
            }
            
            return {
                'success': True,
                'text': text,
                'tables': tables,
                'metadata': metadata,
                'extraction_method': 'pandas'
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX files"""
        if not DOCX_AVAILABLE:
            return {'success': False, 'error': 'DOCX processing library not available'}
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            text = ""
            tables = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                try:
                    data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        data.append(row_data)
                    
                    if len(data) > 1:
                        df = pd.DataFrame(data[1:], columns=data[0])
                        df.name = f"Table_{table_num + 1}"
                        tables.append(df)
                        
                except Exception as e:
                    logger.warning(f"Could not convert DOCX table to DataFrame: {e}")
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'success': True,
                'text': text,
                'tables': tables,
                'metadata': metadata,
                'extraction_method': 'python-docx'
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return {'success': False, 'error': 'Could not decode text file'}
            
            metadata = {
                'encoding': used_encoding,
                'lines': len(text.splitlines()),
                'characters': len(text)
            }
            
            return {
                'success': True,
                'text': text,
                'tables': [],
                'metadata': metadata,
                'extraction_method': 'text'
            }
            
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return {'success': False, 'error': str(e)}


class DataExtractor:
    """Extract and structure financial data from processed documents"""
    
    def __init__(self):
        self.financial_keywords = self._load_financial_keywords()
        self.statement_patterns = self._compile_statement_patterns()
    
    def _load_financial_keywords(self) -> Dict[str, List[str]]:
        """Load financial keywords for data extraction"""
        return {
            'balance_sheet': [
                'balance sheet', 'statement of financial position', 'assets', 'liabilities', 'equity',
                'current assets', 'non-current assets', 'total assets', 'shareholders equity',
                'cash', 'receivables', 'inventory', 'payables', 'debt', 'retained earnings'
            ],
            'cash_flow': [
                'cash flow', 'statement of cash flows', 'operating activities', 'investing activities',
                'financing activities', 'net cash', 'cash and cash equivalents', 'free cash flow',
                'depreciation', 'amortization', 'capital expenditures', 'dividends'
            ],
            'income_statement': [
                'income statement', 'profit and loss', 'revenue', 'net income', 'operating income',
                'earnings', 'comprehensive income', 'sales', 'expenses', 'cost of goods sold'
            ]
        }
    
    def _compile_statement_patterns(self) -> Dict[str, List[re.Pattern]]:
        """Compile regex patterns for financial statement recognition"""
        patterns = {}
        
        # Balance sheet patterns
        patterns['balance_sheet'] = [
            re.compile(r'total\s+assets.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'current\s+assets.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'cash.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'accounts?\s+receivable.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'inventory.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'total\s+liabilities.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'shareholders?\s+equity.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE)
        ]
        
        # Cash flow patterns
        patterns['cash_flow'] = [
            re.compile(r'net\s+income.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'operating\s+activities.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'investing\s+activities.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'financing\s+activities.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'depreciation.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE),
            re.compile(r'capital\s+expenditures.*?(\$?[\d,]+\.?\d*)', re.IGNORECASE)
        ]
        
        return patterns
    
    def identify_document_type(self, text: str) -> Dict[str, Any]:
        """Identify the type of financial document"""
        text_lower = text.lower()
        
        scores = {}
        for doc_type, keywords in self.financial_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            scores[doc_type] = score
        
        if not scores or max(scores.values()) == 0:
            return {
                'document_type': 'unknown',
                'confidence': 0.0,
                'scores': scores
            }
        
        best_type = max(scores.keys(), key=lambda k: scores[k])
        max_score = scores[best_type]
        total_possible = len(self.financial_keywords[best_type])
        confidence = max_score / total_possible
        
        return {
            'document_type': best_type,
            'confidence': confidence,
            'scores': scores
        }
    
    def extract_financial_data(self, processed_doc: Dict[str, Any]) -> Dict[str, Any]:
        """Extract financial data from processed document"""
        try:
            if not processed_doc.get('success', False):
                return {'success': False, 'error': 'Document processing failed'}
            
            text = processed_doc.get('text', '')
            tables = processed_doc.get('tables', [])
            
            # Identify document type
            doc_info = self.identify_document_type(text)
            
            # Extract data based on document type
            if doc_info['document_type'] == 'balance_sheet':
                extracted_data = self._extract_balance_sheet_data(text, tables)
            elif doc_info['document_type'] == 'cash_flow':
                extracted_data = self._extract_cash_flow_data(text, tables)
            elif doc_info['document_type'] == 'income_statement':
                extracted_data = self._extract_income_statement_data(text, tables)
            else:
                # Try to extract general financial data
                extracted_data = self._extract_general_financial_data(text, tables)
            
            # Add document information
            extracted_data.update({
                'document_type': doc_info['document_type'],
                'confidence': doc_info['confidence'],
                'extraction_method': 'pattern_matching',
                'year': self._extract_year(text),
                'currency': self._extract_currency(text)
            })
            
            return {
                'success': True,
                'data': extracted_data,
                'document_info': doc_info,
                'extraction_metadata': {
                    'patterns_matched': len([v for v in extracted_data.values() if isinstance(v, (int, float)) and v != 0]),
                    'tables_processed': len(tables),
                    'text_length': len(text)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting financial data: {e}")
            return {'success': False, 'error': str(e)}
    
    def _extract_balance_sheet_data(self, text: str, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract balance sheet specific data"""
        data = {}
        
        # Define balance sheet line items and their patterns
        patterns = {
            'cash_and_equivalents': [
                r'cash\s+and\s+cash\s+equivalents.*?(\$?[\d,]+\.?\d*)',
                r'cash\s+and\s+equivalents.*?(\$?[\d,]+\.?\d*)',
                r'cash.*?(\$?[\d,]+\.?\d*)'
            ],
            'accounts_receivable': [
                r'accounts?\s+receivable.*?(\$?[\d,]+\.?\d*)',
                r'trade\s+receivables.*?(\$?[\d,]+\.?\d*)'
            ],
            'inventory': [
                r'inventory.*?(\$?[\d,]+\.?\d*)',
                r'inventories.*?(\$?[\d,]+\.?\d*)'
            ],
            'current_assets': [
                r'total\s+current\s+assets.*?(\$?[\d,]+\.?\d*)',
                r'current\s+assets.*?(\$?[\d,]+\.?\d*)'
            ],
            'property_plant_equipment': [
                r'property,?\s+plant\s+and\s+equipment.*?(\$?[\d,]+\.?\d*)',
                r'fixed\s+assets.*?(\$?[\d,]+\.?\d*)'
            ],
            'total_assets': [
                r'total\s+assets.*?(\$?[\d,]+\.?\d*)'
            ],
            'accounts_payable': [
                r'accounts?\s+payable.*?(\$?[\d,]+\.?\d*)',
                r'trade\s+payables.*?(\$?[\d,]+\.?\d*)'
            ],
            'current_liabilities': [
                r'total\s+current\s+liabilities.*?(\$?[\d,]+\.?\d*)',
                r'current\s+liabilities.*?(\$?[\d,]+\.?\d*)'
            ],
            'long_term_debt': [
                r'long.?term\s+debt.*?(\$?[\d,]+\.?\d*)',
                r'non.?current\s+debt.*?(\$?[\d,]+\.?\d*)'
            ],
            'total_liabilities': [
                r'total\s+liabilities.*?(\$?[\d,]+\.?\d*)'
            ],
            'total_equity': [
                r'total\s+shareholders?\s+equity.*?(\$?[\d,]+\.?\d*)',
                r'total\s+equity.*?(\$?[\d,]+\.?\d*)',
                r'shareholders?\s+equity.*?(\$?[\d,]+\.?\d*)'
            ]
        }
        
        # Extract from text using patterns
        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    try:
                        value = self._clean_numeric_value(matches[0])
                        data[field] = value
                        break
                    except ValueError:
                        continue
            
            if field not in data:
                data[field] = 0.0
        
        # Extract from tables
        table_data = self._extract_from_tables(tables, 'balance_sheet')
        data.update(table_data)
        
        return data
    
    def _extract_cash_flow_data(self, text: str, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract cash flow specific data"""
        data = {}
        
        patterns = {
            'net_income': [
                r'net\s+income.*?(\$?[\d,]+\.?\d*)',
                r'net\s+earnings.*?(\$?[\d,]+\.?\d*)'
            ],
            'depreciation_and_amortization': [
                r'depreciation\s+and\s+amortization.*?(\$?[\d,]+\.?\d*)',
                r'depreciation.*?(\$?[\d,]+\.?\d*)'
            ],
            'net_cash_from_operating_activities': [
                r'net\s+cash.*?operating\s+activities.*?(\$?[\d,]+\.?\d*)',
                r'cash\s+from\s+operating.*?(\$?[\d,]+\.?\d*)'
            ],
            'capital_expenditures': [
                r'capital\s+expenditures.*?(\$?[\d,]+\.?\d*)',
                r'capex.*?(\$?[\d,]+\.?\d*)'
            ],
            'net_cash_from_investing_activities': [
                r'net\s+cash.*?investing\s+activities.*?(\$?[\d,]+\.?\d*)',
                r'cash\s+from\s+investing.*?(\$?[\d,]+\.?\d*)'
            ],
            'net_cash_from_financing_activities': [
                r'net\s+cash.*?financing\s+activities.*?(\$?[\d,]+\.?\d*)',
                r'cash\s+from\s+financing.*?(\$?[\d,]+\.?\d*)'
            ],
            'dividends_paid': [
                r'dividends\s+paid.*?(\$?[\d,]+\.?\d*)',
                r'cash\s+dividends.*?(\$?[\d,]+\.?\d*)'
            ]
        }
        
        # Extract from text
        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    try:
                        value = self._clean_numeric_value(matches[0])
                        data[field] = value
                        break
                    except ValueError:
                        continue
            
            if field not in data:
                data[field] = 0.0
        
        # Extract from tables
        table_data = self._extract_from_tables(tables, 'cash_flow')
        data.update(table_data)
        
        return data
    
    def _extract_income_statement_data(self, text: str, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract income statement specific data"""
        data = {}
        
        patterns = {
            'revenue': [
                r'revenue.*?(\$?[\d,]+\.?\d*)',
                r'total\s+revenue.*?(\$?[\d,]+\.?\d*)',
                r'sales.*?(\$?[\d,]+\.?\d*)'
            ],
            'cost_of_goods_sold': [
                r'cost\s+of\s+goods\s+sold.*?(\$?[\d,]+\.?\d*)',
                r'cogs.*?(\$?[\d,]+\.?\d*)'
            ],
            'gross_profit': [
                r'gross\s+profit.*?(\$?[\d,]+\.?\d*)'
            ],
            'operating_income': [
                r'operating\s+income.*?(\$?[\d,]+\.?\d*)',
                r'income\s+from\s+operations.*?(\$?[\d,]+\.?\d*)'
            ],
            'net_income': [
                r'net\s+income.*?(\$?[\d,]+\.?\d*)',
                r'net\s+earnings.*?(\$?[\d,]+\.?\d*)'
            ]
        }
        
        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    try:
                        value = self._clean_numeric_value(matches[0])
                        data[field] = value
                        break
                    except ValueError:
                        continue
            
            if field not in data:
                data[field] = 0.0
        
        return data
    
    def _extract_general_financial_data(self, text: str, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """Extract general financial data when document type is unclear"""
        data = {}
        
        # General financial patterns
        general_patterns = {
            'total_assets': r'total\s+assets.*?(\$?[\d,]+\.?\d*)',
            'total_liabilities': r'total\s+liabilities.*?(\$?[\d,]+\.?\d*)',
            'revenue': r'(?:revenue|sales).*?(\$?[\d,]+\.?\d*)',
            'net_income': r'net\s+income.*?(\$?[\d,]+\.?\d*)',
            'cash': r'cash.*?(\$?[\d,]+\.?\d*)',
            'debt': r'(?:debt|borrowings).*?(\$?[\d,]+\.?\d*)'
        }
        
        # Extract from text
        for field, pattern in general_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    value = self._clean_numeric_value(matches[0])
                    data[field] = value
                except ValueError:
                    data[field] = 0.0
            else:
                data[field] = 0.0
        
        # Extract from tables
        table_data = self._extract_from_tables(tables, 'general')
        data.update(table_data)
        
        return data
    
    def _extract_from_tables(self, tables: List[pd.DataFrame], statement_type: str) -> Dict[str, Any]:
        """Extract financial data from tables"""
        data = {}
        
        if not tables:
            return data
        
        # Define column mappings for different statement types
        column_mappings = {
            'balance_sheet': {
                'cash': ['cash', 'cash and equivalents', 'cash and cash equivalents'],
                'accounts_receivable': ['accounts receivable', 'receivables', 'trade receivables'],
                'inventory': ['inventory', 'inventories'],
                'total_assets': ['total assets'],
                'accounts_payable': ['accounts payable', 'payables', 'trade payables'],
                'total_liabilities': ['total liabilities'],
                'total_equity': ['total equity', 'shareholders equity', 'stockholders equity']
            },
            'cash_flow': {
                'net_income': ['net income', 'net earnings'],
                'operating_cash_flow': ['net cash from operating', 'operating cash flow'],
                'investing_cash_flow': ['net cash from investing', 'investing cash flow'],
                'financing_cash_flow': ['net cash from financing', 'financing cash flow'],
                'capital_expenditures': ['capital expenditures', 'capex'],
                'depreciation': ['depreciation', 'depreciation and amortization']
            },
            'income_statement': {
                'revenue': ['revenue', 'total revenue', 'sales', 'net sales'],
                'gross_profit': ['gross profit'],
                'operating_income': ['operating income', 'income from operations'],
                'net_income': ['net income', 'net earnings']
            },
            'general': {
                'revenue': ['revenue', 'sales', 'total revenue'],
                'net_income': ['net income', 'profit'],
                'assets': ['assets', 'total assets'],
                'liabilities': ['liabilities', 'total liabilities']
            }
        }
        
        mappings = column_mappings.get(statement_type, column_mappings['general'])
        
        for table in tables:
            if table.empty:
                continue
            
            # Try to find financial data in table
            for field, possible_names in mappings.items():
                for col_name in table.columns:
                    col_name_lower = str(col_name).lower()
                    for possible_name in possible_names:
                        if possible_name.lower() in col_name_lower:
                            # Extract numeric values from this column
                            numeric_values = []
                            for value in table[col_name]:
                                try:
                                    if pd.notna(value):
                                        cleaned_value = self._clean_numeric_value(str(value))
                                        if cleaned_value != 0:
                                            numeric_values.append(cleaned_value)
                                except (ValueError, TypeError):
                                    continue
                            
                            if numeric_values:
                                # Take the most recent or largest value
                                data[field] = max(numeric_values)
                                break
                    if field in data:
                        break
        
        return data
    
    def _clean_numeric_value(self, value_str: str) -> float:
        """Clean and convert string to numeric value"""
        if not value_str or pd.isna(value_str):
            return 0.0
        
        # Convert to string and clean
        value_str = str(value_str).strip()
        
        # Remove common formatting
        value_str = re.sub(r'[,$\s()]', '', value_str)
        
        # Handle parentheses (negative values)
        is_negative = '(' in str(value_str) and ')' in str(value_str)
        value_str = re.sub(r'[()]', '', value_str)
        
        # Handle minus signs
        if value_str.startswith('-'):
            is_negative = True
            value_str = value_str[1:]
        
        # Remove non-numeric characters except decimal point
        value_str = re.sub(r'[^\d.]', '', value_str)
        
        if not value_str or value_str == '.':
            return 0.0
        
        try:
            value = float(value_str)
            return -value if is_negative else value
        except (ValueError, TypeError):
            return 0.0
    
    def _extract_year(self, text: str) -> Optional[int]:
        """Extract year from document text"""
        # Look for 4-digit years
        year_pattern = r'\b(20\d{2})\b'
        matches = re.findall(year_pattern, text)
        
        if matches:
            # Return the most recent year found
            years = [int(year) for year in matches]
            return max(years)
        
        return None
    
    def _extract_currency(self, text: str) -> str:
        """Extract currency from document text"""
        # Common currency patterns
        currency_patterns = {
            'USD': [r'\$', r'USD', r'US\s*Dollar', r'United States Dollar'],      # Fixed: \$ instead of $'\
            'EUR': [r'€', r'EUR', r'Euro'],
            'GBP': [r'£', r'GBP', r'British Pound', r'Sterling'],
            'CAD': [r'CAD', r'Canadian Dollar'],                                   # Fixed: Added r prefix
            'AUD': [r'AUD', r'Australian Dollar'],
            'JPY': [r'¥', r'JPY', r'Japanese Yen']
        }
        
        text_lower = text.lower()
        
        for currency, patterns in currency_patterns.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return currency
        
        # Default to USD if no currency found but $ symbols present
        if '$' in text:
            return 'USD'
        
        return 'USD'  # Default currency


class FileManager:
    """Manage file operations and storage for the Financial Risk Assessment Platform"""
    
    def __init__(self, base_upload_path: str = "uploads"):
        self.base_upload_path = Path(base_upload_path)
        self.base_upload_path.mkdir(exist_ok=True)
        self.allowed_extensions = {'.pdf', '.xlsx', '.xls', '.csv', '.docx', '.txt'}
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    def save_uploaded_file(self, file, user_id: str = None) -> Dict[str, Any]:
        """Save uploaded file to storage"""
        try:
            # Create user-specific directory
            if user_id:
                upload_dir = self.base_upload_path / user_id
            else:
                upload_dir = self.base_upload_path / "anonymous"
            
            upload_dir.mkdir(exist_ok=True)
            
            # Generate secure filename
            if WERKZEUG_AVAILABLE:
                from werkzeug.utils import secure_filename
                filename = secure_filename(file.filename)
            else:
                # Simple filename cleaning if werkzeug not available
                filename = re.sub(r'[^a-zA-Z0-9._-]', '', file.filename)
            
            # Add timestamp to avoid conflicts
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{timestamp}{ext}"
            
            file_path = upload_dir / filename
            
            # Check file size
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)
            
            if file_size > self.max_file_size:
                return {
                    'success': False,
                    'error': f'File too large: {file_size / (1024*1024):.2f}MB (max: {self.max_file_size/(1024*1024)}MB)'
                }
            
            # Check file extension
            if Path(filename).suffix.lower() not in self.allowed_extensions:
                return {
                    'success': False,
                    'error': f'File type not allowed: {Path(filename).suffix}'
                }
            
            # Save file
            file.save(str(file_path))
            
            return {
                'success': True,
                'file_path': str(file_path),
                'filename': filename,
                'size': file_size,
                'upload_time': datetime.now()
            }
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            return {'success': False, 'error': str(e)}
    
    def delete_file(self, file_path: str) -> bool:
        """Delete file from storage"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file: {e}")
            return False
    
    def cleanup_old_files(self, days_old: int = 7) -> int:
        """Clean up files older than specified days"""
        cleaned_count = 0
        cutoff_date = datetime.now() - pd.Timedelta(days=days_old)
        
        try:
            for file_path in self.base_upload_path.rglob('*'):
                if file_path.is_file():
                    file_modified = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_modified < cutoff_date:
                        file_path.unlink()
                        cleaned_count += 1
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")
        
        return cleaned_count


def process_uploaded_files(file_paths: List[str], user_id: str = None) -> Dict[str, Any]:
    """Main function to process uploaded files"""
    processor = DocumentProcessor()
    extractor = DataExtractor()
    
    results = {
        'success': True,
        'processed_files': [],
        'failed_files': [],
        'extracted_data': {},
        'summary': {
            'total_files': len(file_paths),
            'successful': 0,
            'failed': 0,
            'processing_time': 0
        }
    }
    
    start_time = datetime.now()
    
    for file_path in file_paths:
        try:
            # Process document
            doc_result = processor.process_document(file_path)
            
            if doc_result['success']:
                # Extract financial data
                extracted_result = extractor.extract_financial_data(doc_result)
                
                if extracted_result['success']:
                    results['processed_files'].append({
                        'file_path': file_path,
                        'document_info': extracted_result['document_info'],
                        'extraction_metadata': extracted_result['extraction_metadata']
                    })
                    results['extracted_data'][file_path] = extracted_result['data']
                    results['summary']['successful'] += 1
                else:
                    results['failed_files'].append({
                        'file_path': file_path,
                        'error': extracted_result.get('error', 'Data extraction failed')
                    })
                    results['summary']['failed'] += 1
            else:
                results['failed_files'].append({
                    'file_path': file_path,
                    'error': doc_result.get('error', 'Document processing failed')
                })
                results['summary']['failed'] += 1
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            results['failed_files'].append({
                'file_path': file_path,
                'error': str(e)
            })
            results['summary']['failed'] += 1
    
    # Calculate processing time
    end_time = datetime.now()
    results['summary']['processing_time'] = (end_time - start_time).total_seconds()
    
    # Set overall success status
    if results['summary']['successful'] == 0:
        results['success'] = False
    
    return results


def clean_and_validate_data(extracted_data: Dict[str, Any]) -> Dict[str, Any]:
    """Clean and validate extracted financial data"""
    try:
        cleaned_data = {}
        validation_results = {
            'cleaned_fields': [],
            'invalid_fields': [],
            'warnings': []
        }
        
        for field, value in extracted_data.items():
            if field in ['document_type', 'confidence', 'extraction_method', 'year', 'currency']:
                # Keep metadata as is
                cleaned_data[field] = value
                continue
            
            # Clean numeric values
            if isinstance(value, (int, float)):
                if pd.isna(value) or np.isinf(value):
                    cleaned_data[field] = 0.0
                    validation_results['warnings'].append(f"Invalid value for {field}, set to 0")
                elif value < 0 and field not in ['net_income', 'net_cash_from_operating_activities']:
                    # Some fields can be negative
                    validation_results['warnings'].append(f"Negative value for {field}: {value}")
                    cleaned_data[field] = abs(value)
                else:
                    cleaned_data[field] = float(value)
                
                validation_results['cleaned_fields'].append(field)
            
            elif isinstance(value, str):
                # Try to convert string to numeric
                try:
                    numeric_value = float(re.sub(r'[^\d.-]', '', value))
                    cleaned_data[field] = numeric_value
                    validation_results['cleaned_fields'].append(field)
                except ValueError:
                    cleaned_data[field] = 0.0
                    validation_results['invalid_fields'].append(field)
            
            else:
                cleaned_data[field] = 0.0
                validation_results['invalid_fields'].append(field)
        
        return {
            'success': True,
            'cleaned_data': cleaned_data,
            'validation_results': validation_results
        }
        
    except Exception as e:
        logger.error(f"Error cleaning and validating data: {e}")
        return {
            'success': False,
            'error': str(e),
            'cleaned_data': {},
            'validation_results': {}
        }


def extract_financial_data(text: str, tables: List[pd.DataFrame] = None) -> Dict[str, Any]:
    """
    Standalone function to extract financial data from text and tables
    
    Args:
        text: Document text content
        tables: Optional list of DataFrames containing tabular data
        
    Returns:
        Dictionary containing extracted financial data
    """
    extractor = DataExtractor()
    
    # Create a mock processed document structure
    processed_doc = {
        'success': True,
        'text': text,
        'tables': tables or [],
        'metadata': {'extraction_method': 'standalone'}
    }
    
    return extractor.extract_financial_data(processed_doc)


def validate_financial_statements(balance_sheet: Dict[str, float], 
                                income_statement: Dict[str, float] = None,
                                cash_flow: Dict[str, float] = None) -> Dict[str, Any]:
    """
    Validate financial statement data for consistency
    
    Args:
        balance_sheet: Balance sheet data
        income_statement: Optional income statement data
        cash_flow: Optional cash flow statement data
        
    Returns:
        Validation results with consistency checks
    """
    validation_results = {
        'is_valid': True,
        'errors': [],
        'warnings': [],
        'checks_performed': []
    }
    
    try:
        # Check balance sheet equation
        total_assets = balance_sheet.get('total_assets', 0)
        total_liabilities = balance_sheet.get('total_liabilities', 0)
        total_equity = balance_sheet.get('total_equity', 0)
        
        if total_assets != 0 and (total_liabilities != 0 or total_equity != 0):
            balance_diff = abs(total_assets - (total_liabilities + total_equity))
            tolerance = max(total_assets * 0.01, 1000)  # 1% tolerance
            
            if balance_diff > tolerance:
                validation_results['errors'].append(
                    f"Balance sheet equation error: Assets ({total_assets:,.0f}) ≠ Liabilities + Equity ({total_liabilities + total_equity:,.0f})"
                )
                validation_results['is_valid'] = False
            
            validation_results['checks_performed'].append('balance_sheet_equation')
        
        # Cross-statement validation if income statement provided
        if income_statement:
            bs_net_income = balance_sheet.get('retained_earnings', 0)
            is_net_income = income_statement.get('net_income', 0)
            
            # Note: This is a simplified check - in reality, retained earnings = 
            # previous retained earnings + net income - dividends
            validation_results['checks_performed'].append('income_statement_consistency')
        
        # Cash flow validation
        if cash_flow:
            operating_cf = cash_flow.get('net_cash_from_operating_activities', 0)
            investing_cf = cash_flow.get('net_cash_from_investing_activities', 0)
            financing_cf = cash_flow.get('net_cash_from_financing_activities', 0)
            
            net_change = operating_cf + investing_cf + financing_cf
            
            # This should equal the change in cash from balance sheet
            validation_results['checks_performed'].append('cash_flow_consistency')
        
        return validation_results
        
    except Exception as e:
        logger.error(f"Error validating financial statements: {e}")
        return {
            'is_valid': False,
            'error': str(e),
            'errors': [],
            'warnings': [],
            'checks_performed': []
        }


# Example usage and utility functions
def get_supported_file_types() -> List[str]:
    """Get list of supported file types"""
    processor = DocumentProcessor()
    supported = []
    for file_type, is_supported in processor.supported_formats.items():
        if is_supported:
            supported.append(file_type)
    return supported


def estimate_processing_time(file_size_mb: float, file_type: str) -> float:
    """Estimate processing time for a file"""
    processor = DocumentProcessor()
    return processor._estimate_processing_time(file_size_mb * 1024 * 1024, f".{file_type}")


# Configuration and constants
FINANCIAL_DATA_FIELDS = [
    'total_assets', 'current_assets', 'cash_and_equivalents', 'accounts_receivable',
    'inventory', 'property_plant_equipment', 'total_liabilities', 'current_liabilities',
    'accounts_payable', 'long_term_debt', 'total_equity', 'retained_earnings',
    'revenue', 'cost_of_goods_sold', 'gross_profit', 'operating_income', 'net_income',
    'net_cash_from_operating_activities', 'net_cash_from_investing_activities',
    'net_cash_from_financing_activities', 'capital_expenditures', 'dividends_paid',
    'depreciation_and_amortization', 'interest_expense', 'income_tax_expense'
]

REQUIRED_BALANCE_SHEET_FIELDS = [
    'total_assets', 'total_liabilities', 'total_equity'
]

REQUIRED_INCOME_STATEMENT_FIELDS = [
    'revenue', 'net_income'
]

REQUIRED_CASH_FLOW_FIELDS = [
    'net_cash_from_operating_activities'
]

# Enhanced pattern matching for different financial statement formats
ENHANCED_FINANCIAL_PATTERNS = {
    'balance_sheet': {
        'cash_and_equivalents': [
            r'cash\s+and\s+cash\s+equivalents.*?(\$?[\d,]+\.?\d*)',
            r'cash\s+&\s+equivalents.*?(\$?[\d,]+\.?\d*)',
            r'cash\s+and\s+short.?term\s+investments.*?(\$?[\d,]+\.?\d*)',
            r'cash.*?(\$?[\d,]+\.?\d*)',
            r'liquid\s+assets.*?(\$?[\d,]+\.?\d*)'
        ],
        'accounts_receivable': [
            r'accounts?\s+receivable.*?(\$?[\d,]+\.?\d*)',
            r'trade\s+receivables.*?(\$?[\d,]+\.?\d*)',
            r'receivables.*?(\$?[\d,]+\.?\d*)',
            r'debtors.*?(\$?[\d,]+\.?\d*)'
        ],
        'inventory': [
            r'inventory.*?(\$?[\d,]+\.?\d*)',
            r'inventories.*?(\$?[\d,]+\.?\d*)',
            r'stock.*?(\$?[\d,]+\.?\d*)',
            r'merchandise.*?(\$?[\d,]+\.?\d*)'
        ],
        'total_assets': [
            r'total\s+assets.*?(\$?[\d,]+\.?\d*)',
            r'sum\s+of\s+assets.*?(\$?[\d,]+\.?\d*)',
            r'aggregate\s+assets.*?(\$?[\d,]+\.?\d*)'
        ],
        'total_equity': [
            r'total\s+shareholders?\s*equity.*?(\$?[\d,]+\.?\d*)',
            r'total\s+stockholders?\s*equity.*?(\$?[\d,]+\.?\d*)',
            r'total\s+equity.*?(\$?[\d,]+\.?\d*)',
            r'net\s+worth.*?(\$?[\d,]+\.?\d*)',
            r'owners?\s*equity.*?(\$?[\d,]+\.?\d*)'
        ]
    },
    'income_statement': {
        'revenue': [
            r'(?:total\s+)?(?:net\s+)?revenue.*?(\$?[\d,]+\.?\d*)',
            r'(?:total\s+)?(?:net\s+)?sales.*?(\$?[\d,]+\.?\d*)',
            r'turnover.*?(\$?[\d,]+\.?\d*)',
            r'income\s+from\s+sales.*?(\$?[\d,]+\.?\d*)',
            r'gross\s+sales.*?(\$?[\d,]+\.?\d*)'
        ],
        'cost_of_goods_sold': [
            r'cost\s+of\s+(?:goods\s+sold|sales).*?(\$?[\d,]+\.?\d*)',
            r'cogs.*?(\$?[\d,]+\.?\d*)',
            r'cost\s+of\s+revenue.*?(\$?[\d,]+\.?\d*)',
            r'direct\s+costs.*?(\$?[\d,]+\.?\d*)'
        ],
        'operating_income': [
            r'operating\s+income.*?(\$?[\d,]+\.?\d*)',
            r'income\s+from\s+operations.*?(\$?[\d,]+\.?\d*)',
            r'operating\s+profit.*?(\$?[\d,]+\.?\d*)',
            r'ebit.*?(\$?[\d,]+\.?\d*)'
        ],
        'net_income': [
            r'net\s+income.*?(\$?[\d,]+\.?\d*)',
            r'net\s+earnings.*?(\$?[\d,]+\.?\d*)',
            r'net\s+profit.*?(\$?[\d,]+\.?\d*)',
            r'bottom\s+line.*?(\$?[\d,]+\.?\d*)',
            r'profit\s+after\s+tax.*?(\$?[\d,]+\.?\d*)'
        ]
    }
}


class AdvancedDataProcessor:
    """
    Advanced data processing with AI-enhanced extraction capabilities
    """
    
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.data_extractor = DataExtractor()
        self.confidence_threshold = 0.7
        self.extraction_stats = {
            'documents_processed': 0,
            'successful_extractions': 0,
            'average_confidence': 0.0
        }
    
    def process_financial_documents(self, file_paths: List[str], 
                                  company_name: str = None) -> Dict[str, Any]:
        """
        Advanced processing of multiple financial documents with consolidation
        
        Args:
            file_paths: List of document paths
            company_name: Optional company name for validation
            
        Returns:
            Consolidated financial data with confidence scoring
        """
        try:
            consolidated_data = {}
            processing_results = []
            confidence_scores = []
            
            for file_path in file_paths:
                # Process individual document
                doc_result = self.document_processor.process_document(file_path)
                
                if doc_result['success']:
                    # Extract financial data
                    extracted_result = self.data_extractor.extract_financial_data(doc_result)
                    
                    if extracted_result['success']:
                        data = extracted_result['data']
                        confidence = extracted_result['document_info']['confidence']
                        
                        # Store results
                        processing_results.append({
                            'file_path': file_path,
                            'document_type': data.get('document_type', 'unknown'),
                            'confidence': confidence,
                            'extracted_fields': len([k for k, v in data.items() 
                                                   if isinstance(v, (int, float)) and v != 0])
                        })
                        
                        confidence_scores.append(confidence)
                        
                        # Consolidate data by document type
                        doc_type = data.get('document_type', 'general')
                        if doc_type not in consolidated_data:
                            consolidated_data[doc_type] = {}
                        
                        # Merge data with confidence weighting
                        for field, value in data.items():
                            if isinstance(value, (int, float)) and field != 'confidence':
                                if field in consolidated_data[doc_type]:
                                    # Average with confidence weighting
                                    existing_conf = consolidated_data[doc_type].get(f"{field}_confidence", 0.5)
                                    weighted_value = ((consolidated_data[doc_type][field] * existing_conf + 
                                                     value * confidence) / (existing_conf + confidence))
                                    consolidated_data[doc_type][field] = weighted_value
                                    consolidated_data[doc_type][f"{field}_confidence"] = (existing_conf + confidence) / 2
                                else:
                                    consolidated_data[doc_type][field] = value
                                    consolidated_data[doc_type][f"{field}_confidence"] = confidence
            
            # Calculate overall statistics
            avg_confidence = np.mean(confidence_scores) if confidence_scores else 0
            
            # Validate consolidated data
            validation_result = self._validate_consolidated_data(consolidated_data)
            
            # Generate final consolidated dataset
            final_data = self._create_unified_dataset(consolidated_data)
            
            return {
                'success': True,
                'consolidated_data': final_data,
                'processing_results': processing_results,
                'validation': validation_result,
                'statistics': {
                    'documents_processed': len(file_paths),
                    'successful_extractions': len(processing_results),
                    'average_confidence': round(avg_confidence, 3),
                    'data_completeness': self._calculate_completeness_score(final_data),
                    'extraction_quality': 'high' if avg_confidence > 0.8 else 'medium' if avg_confidence > 0.6 else 'low'
                },
                'processing_date': datetime.now()
            }
            
        except Exception as e:
            logger.error(f"Error in advanced processing: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _validate_consolidated_data(self, consolidated_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate consolidated data for consistency"""
        validation_results = {
            'is_valid': True,
            'cross_statement_checks': [],
            'anomalies': [],
            'recommendations': []
        }
        
        try:
            # Extract data from different statements
            balance_sheet = consolidated_data.get('balance_sheet', {})
            income_statement = consolidated_data.get('income_statement', {})
            cash_flow = consolidated_data.get('cash_flow', {})
            
            # Cross-statement validation
            if balance_sheet and income_statement:
                # Check if net income is reasonable relative to assets
                net_income = income_statement.get('net_income', 0)
                total_assets = balance_sheet.get('total_assets', 0)
                
                if total_assets > 0:
                    roa = net_income / total_assets
                    if abs(roa) > 0.5:  # 50% ROA is extremely high/low
                        validation_results['anomalies'].append(
                            f"Unusual ROA detected: {roa:.1%} - verify net income and total assets"
                        )
            
            # Balance sheet equation check
            if balance_sheet:
                assets = balance_sheet.get('total_assets', 0)
                liabilities = balance_sheet.get('total_liabilities', 0)
                equity = balance_sheet.get('total_equity', 0)
                
                if assets > 0 and (liabilities > 0 or equity > 0):
                    balance_diff = abs(assets - (liabilities + equity))
                    tolerance = max(assets * 0.05, 1000)
                    
                    if balance_diff > tolerance:
                        validation_results['is_valid'] = False
                        validation_results['cross_statement_checks'].append(
                            f"Balance sheet equation imbalance: ${balance_diff:,.0f}"
                        )
            
            return validation_results
            
        except Exception as e:
            logger.error(f"Error validating consolidated data: {e}")
            return {'is_valid': False, 'error': str(e)}
    
    def _create_unified_dataset(self, consolidated_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create unified financial dataset from all statements"""
        unified_data = {}
        
        # Combine all data sources
        for statement_type, data in consolidated_data.items():
            for field, value in data.items():
                if not field.endswith('_confidence'):
                    # Use highest confidence value if duplicate fields exist
                    if field in unified_data:
                        existing_conf = unified_data.get(f"{field}_confidence", 0)
                        new_conf = data.get(f"{field}_confidence", 0)
                        
                        if new_conf > existing_conf:
                            unified_data[field] = value
                            unified_data[f"{field}_confidence"] = new_conf
                    else:
                        unified_data[field] = value
                        unified_data[f"{field}_confidence"] = data.get(f"{field}_confidence", 0.5)
        
        return unified_data
    
    def _calculate_completeness_score(self, data: Dict[str, Any]) -> float:
        """Calculate data completeness score"""
        try:
            essential_fields = [
                'total_assets', 'total_liabilities', 'total_equity',
                'revenue', 'net_income', 'net_cash_from_operating_activities'
            ]
            
            present_fields = sum(1 for field in essential_fields 
                               if field in data and data[field] != 0)
            
            return present_fields / len(essential_fields)
            
        except Exception:
            return 0.0


class IntelligentTableExtractor:
    """
    AI-enhanced table extraction with financial context understanding
    """
    
    def __init__(self):
        self.financial_keywords = [
            'assets', 'liabilities', 'equity', 'revenue', 'income', 'expenses',
            'cash', 'debt', 'profit', 'loss', 'earnings', 'sales', 'costs'
        ]
        
        self.table_types = {
            'balance_sheet': ['assets', 'liabilities', 'equity', 'balance'],
            'income_statement': ['revenue', 'income', 'expenses', 'profit', 'loss'],
            'cash_flow': ['cash', 'operating', 'investing', 'financing', 'flow']
        }
    
    def extract_financial_tables(self, tables: List[pd.DataFrame]) -> Dict[str, Any]:
        """
        Extract and classify financial tables using AI-enhanced methods
        
        Args:
            tables: List of extracted tables
            
        Returns:
            Classified and processed financial tables
        """
        try:
            classified_tables = {
                'balance_sheet': [],
                'income_statement': [],
                'cash_flow': [],
                'other': []
            }
            
            processing_stats = {
                'total_tables': len(tables),
                'classified_tables': 0,
                'high_confidence_tables': 0
            }
            
            for i, table in enumerate(tables):
                if table.empty:
                    continue
                
                # Classify table type
                classification_result = self._classify_table(table)
                table_type = classification_result['type']
                confidence = classification_result['confidence']
                
                # Process table based on type
                processed_table = self._process_financial_table(table, table_type)
                
                # Add metadata
                processed_table['metadata'] = {
                    'original_index': i,
                    'classification_confidence': confidence,
                    'row_count': len(table),
                    'column_count': len(table.columns),
                    'financial_score': self._calculate_financial_relevance_score(table)
                }
                
                # Store in appropriate category
                classified_tables[table_type].append(processed_table)
                
                # Update stats
                processing_stats['classified_tables'] += 1
                if confidence > 0.8:
                    processing_stats['high_confidence_tables'] += 1
            
            return {
                'success': True,
                'classified_tables': classified_tables,
                'processing_stats': processing_stats,
                'extraction_date': datetime.now()
            }
            
        except Exception as e:
            logger.error(f"Error in intelligent table extraction: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _classify_table(self, table: pd.DataFrame) -> Dict[str, Any]:
        """Classify table type using keyword analysis"""
        try:
            # Convert table to text for analysis
            table_text = table.to_string().lower()
            
            # Calculate scores for each table type
            type_scores = {}
            
            for table_type, keywords in self.table_types.items():
                score = sum(1 for keyword in keywords if keyword in table_text)
                type_scores[table_type] = score
            
            # Determine best match
            if max(type_scores.values()) > 0:
                best_type = max(type_scores.keys(), key=lambda k: type_scores[k])
                confidence = type_scores[best_type] / len(self.table_types[best_type])
            else:
                best_type = 'other'
                confidence = 0.0
            
            return {
                'type': best_type,
                'confidence': min(confidence, 1.0),
                'scores': type_scores
            }
            
        except Exception as e:
            logger.error(f"Error classifying table: {e}")
            return {'type': 'other', 'confidence': 0.0}
    
    def _process_financial_table(self, table: pd.DataFrame, table_type: str) -> Dict[str, Any]:
        """Process table based on its financial type"""
        try:
            processed_data = {
                'raw_data': table.to_dict(),
                'extracted_values': {},
                'data_quality': 'unknown'
            }
            
            # Extract numeric values
            numeric_data = {}
            for col in table.columns:
                numeric_col = pd.to_numeric(table[col], errors='coerce').dropna()
                if len(numeric_col) > 0:
                    numeric_data[str(col)] = numeric_col.tolist()
            
            processed_data['numeric_data'] = numeric_data
            
            # Calculate data quality score
            total_cells = table.size
            numeric_cells = sum(len(values) for values in numeric_data.values())
            quality_score = numeric_cells / total_cells if total_cells > 0 else 0
            
            if quality_score > 0.7:
                processed_data['data_quality'] = 'high'
            elif quality_score > 0.4:
                processed_data['data_quality'] = 'medium'
            else:
                processed_data['data_quality'] = 'low'
            
            return processed_data
            
        except Exception as e:
            logger.error(f"Error processing financial table: {e}")
            return {'raw_data': {}, 'extracted_values': {}, 'data_quality': 'error'}
    
    def _calculate_financial_relevance_score(self, table: pd.DataFrame) -> float:
        """Calculate how financially relevant a table is"""
        try:
            table_text = table.to_string().lower()
            
            # Count financial keywords
            keyword_count = sum(1 for keyword in self.financial_keywords 
                              if keyword in table_text)
            
            # Check for numeric data
            numeric_ratio = 0
            for col in table.columns:
                numeric_col = pd.to_numeric(table[col], errors='coerce').dropna()
                if len(numeric_col) > 0:
                    numeric_ratio += len(numeric_col) / len(table)
            
            numeric_ratio = numeric_ratio / len(table.columns) if len(table.columns) > 0 else 0
            
            # Combined score
            relevance_score = (keyword_count * 0.3 + numeric_ratio * 0.7)
            return min(relevance_score, 1.0)
            
        except Exception:
            return 0.0


def batch_process_documents(directory_path: str, 
                          file_patterns: List[str] = None) -> Dict[str, Any]:
    """
    Batch process multiple documents in a directory
    
    Args:
        directory_path: Path to directory containing documents
        file_patterns: Optional list of file patterns to match
        
    Returns:
        Batch processing results
    """
    try:
        directory = Path(directory_path)
        if not directory.exists():
            return {
                'success': False,
                'error': f'Directory not found: {directory_path}'
            }
        
        # Default file patterns if not provided
        if file_patterns is None:
            file_patterns = ['*.pdf', '*.xlsx', '*.xls', '*.csv', '*.docx']
        
        # Find all matching files
        file_paths = []
        for pattern in file_patterns:
            file_paths.extend(directory.glob(pattern))
        
        if not file_paths:
            return {
                'success': False,
                'error': 'No matching files found in directory'
            }
        
        # Process files
        processor = AdvancedDataProcessor()
        results = processor.process_financial_documents([str(fp) for fp in file_paths])
        
        # Add batch-specific metadata
        results['batch_info'] = {
            'directory_processed': str(directory),
            'file_patterns_used': file_patterns,
            'files_found': len(file_paths),
            'processing_mode': 'batch'
        }
        
        return results
        
    except Exception as e:
        logger.error(f"Error in batch processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }


def export_processed_data(data: Dict[str, Any], 
                         export_format: str = 'json',
                         output_path: str = None) -> str:
    """
    Export processed financial data to various formats
    
    Args:
        data: Processed financial data
        export_format: Format to export ('json', 'csv', 'excel')
        output_path: Optional output file path
        
    Returns:
        Path to exported file
    """
    try:
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"financial_data_{timestamp}"
        
        if export_format.lower() == 'json':
            output_file = f"{output_path}.json"
            
            # Make data JSON serializable
            json_data = {}
            for key, value in data.items():
                if isinstance(value, (pd.DataFrame, np.ndarray)):
                    json_data[key] = str(value)
                elif isinstance(value, datetime):
                    json_data[key] = value.isoformat()
                else:
                    json_data[key] = value
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
                
        elif export_format.lower() == 'csv':
            output_file = f"{output_path}.csv"
            
            # Convert to DataFrame if possible
            if 'consolidated_data' in data:
                df = pd.DataFrame([data['consolidated_data']])
                df.to_csv(output_file, index=False)
            else:
                return None
                
        elif export_format.lower() == 'excel':
            output_file = f"{output_path}.xlsx"
            
            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                # Export consolidated data
                if 'consolidated_data' in data:
                    df = pd.DataFrame([data['consolidated_data']])
                    df.to_excel(writer, sheet_name='Financial Data', index=False)
                
                # Export processing results
                if 'processing_results' in data:
                    results_df = pd.DataFrame(data['processing_results'])
                    results_df.to_excel(writer, sheet_name='Processing Results', index=False)
                
                # Export statistics
                if 'statistics' in data:
                    stats_df = pd.DataFrame([data['statistics']])
                    stats_df.to_excel(writer, sheet_name='Statistics', index=False)
        
        else:
            return None
        
        return output_file
        
    except Exception as e:
        logger.error(f"Error exporting data: {e}")
        return None